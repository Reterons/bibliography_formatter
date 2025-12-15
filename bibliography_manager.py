import re
import json
from typing import List, Optional, Tuple, Dict, Any
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from bibliography import *
from citation_style import CitationStyle
from author import Author

class BibliographyManager:
    def __init__(self):
        self.items: List[BibliographicItem] = []
        self.current_style: Optional[CitationStyle] = None
        self.available_fields = [
            'authors_str', 'title', 'year', 'journal', 'volume',
            'issue', 'pages', 'publisher', 'edition', 'isbn', 'city',
            'conference_name', 'location', 'url', 'doi', 'accessed_date',
            'website'
        ]

    def add_item(self, item):
        self.items.append(item)

    def _clean_extracted_value(self, value, field_name=""):
        if value is None:
            return ""
        v = str(value).strip()
        if not v:
            return ""

        if field_name:
            v = re.sub(rf'^\s*{re.escape(field_name)}\s*=\s*', '', v, flags=re.IGNORECASE)

        while v.startswith('='):
            v = v[1:].lstrip()

        if len(v) >= 2 and ((v[0] == '(' and v[-1] == ')') or (v[0] == '[' and v[-1] == ']') or (v[0] == '{' and v[-1] == '}')):
            v = v[1:-1].strip()
            while v.startswith('='):
                v = v[1:].lstrip()

        v = v.rstrip('.,;:!?)\"]}>')
        v = v.rstrip()
        return v

    def _extract_hyperlinks_from_paragraph(self, para):
        urls = []
        try:
            for h in para._p.xpath('.//w:hyperlink'):
                rid = h.get(qn('r:id'))
                if not rid:
                    continue
                rel = para.part.rels.get(rid)
                if not rel:
                    continue
                target = getattr(rel, 'target_ref', None) or getattr(rel, 'target', None)
                if not target:
                    continue
                urls.append(str(target))
        except Exception:
            return []

        seen = set()
        out = []
        for u in urls:
            if u not in seen:
                seen.add(u)
                out.append(u)
        return out

    def parse_docx(self, filepath):
        items = []
        doc = Document(filepath)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            urls = self._extract_hyperlinks_from_paragraph(para)

            raw_lines = [s.strip() for s in re.split(r'\n\s*\d+\.\s*', text) if s.strip()]
            if len(raw_lines) == 1:
                raw_lines = [re.sub(r'^\s*\d+\.\s*', '', raw_lines[0]).strip()]

            if urls:
                if len(urls) == len(raw_lines):
                    raw_lines = [
                        (ln if self._extract_url(ln) else f"{ln} {urls[i]}")
                        for i, ln in enumerate(raw_lines)
                    ]
                else:
                    last = raw_lines[-1]
                    if not self._extract_url(last):
                        raw_lines[-1] = f"{last} {urls[-1]}"

            for line in raw_lines:
                line = line.strip()
                if not line:
                    continue
                item = self._parse_reference(line)
                if item:
                    items.append(item)

        return items

    def _parse_reference(self, text):
        text = text.strip()
        if not text:
            return None

        text = re.sub(r'\s+', ' ', text)
        text_lower = text.lower()

        if (any(x in text_lower for x in ['vol.', 'no.', 'pp.', 'p. ', 'journal', 'trans.', 'proc.']) or
            'doi:' in text_lower or
            re.search(r'\d{4},\s*\d+,\s*\d+[-–]\d+', text)):
            return self._parse_article(text)

        if any(x in text_lower for x in ['ed.', 'edition', 'publisher', 'press', 'isbn', 'chap.']):
            return self._parse_book(text)

        if any(x in text_lower for x in ['conference', 'proc.', 'proceedings', 'symposium', 'workshop']):
            return self._parse_conference(text)

        if any(x in text_lower for x in ['http://', 'https://', 'www.', 'accessed', 'retrieved']):
            return self._parse_electronic_resource(text)

        return self._parse_article(text)

    def _extract_authors_and_title(self, text):
        text = text.strip()

        quote_patterns = [
            r'["\u201c](.*?)["\u201d]',
            r'["\'](.*?)["\']',
        ]

        for pattern in quote_patterns:
            title_match = re.search(pattern, text)
            if title_match:
                title = title_match.group(1).strip()
                title = re.sub(r'[.,;:\s]+$', '', title)
                authors_text = text[:title_match.start()].strip()
                authors_text = authors_text.rstrip(',')
                rest = text[title_match.end():].strip()
                authors = self._parse_authors(authors_text)
                return authors, title, rest

        if ';' in text:
            pattern = r'^(.*?[A-Za-z]+,\s*[A-Z]\.(?:\.[A-Z])?\.)\s+(.*)$'
            match = re.match(pattern, text)

            if match:
                authors_text = match.group(1).strip()
                rest_after_authors = match.group(2).strip()

                title_patterns = [
                    r'^(.*?)(?:\.\s+[A-Z]\.)',
                    r'^(.*?)(?:\.\s+[A-Z][a-z])',
                    r'^(.*?)(?:\.\s+\d{4})',
                ]

                for title_pattern in title_patterns:
                    title_match = re.search(title_pattern, rest_after_authors)
                    if title_match:
                        title = title_match.group(1).strip()
                        rest = rest_after_authors[title_match.end():].strip()
                        authors = self._parse_authors(authors_text)
                        return authors, title, rest

                title = rest_after_authors
                rest = ""
                authors = self._parse_authors(authors_text)
                return authors, title, rest

        pattern = r'^(.*?\.)\s+([A-Z].*?)(?=\.\s+[A-Z]|\.\s*\d{4}|,\s*\d{4}|$)'
        match = re.search(pattern, text, re.DOTALL)

        if match:
            authors_text = match.group(1).strip()
            title = match.group(2).strip()
            title_end = match.end(2)
            rest = text[title_end:].strip()

            if rest.startswith('.'):
                rest = rest[1:].strip()

            authors = self._parse_authors(authors_text)
            return authors, title, rest

        return [], "", text

    def _parse_authors(self, author_text):
        authors = []
        if not author_text:
            return authors

        author_text = re.sub(r'\s+', ' ', author_text.strip())

        if ';' in author_text:
            parts = [p.strip() for p in author_text.split(';')]
            for part in parts:
                if part:
                    author = self._parse_single_author(part)
                    if author and author.last_name:
                        authors.append(author)
            return authors

        if ',' in author_text:
            parts = [p.strip() for p in author_text.split(',') if p.strip()]

            if len(parts) == 2:
                author = self._parse_single_author(f"{parts[0]}, {parts[1]}")
                if author and author.last_name:
                    return [author]

            if len(parts) % 2 == 0 and len(parts) > 2:
                for i in range(0, len(parts), 2):
                    author = self._parse_single_author(f"{parts[i]}, {parts[i+1]}")
                    if author and author.last_name:
                        authors.append(author)
                if authors:
                    return authors

            for part in parts:
                if part and not re.match(r'^[A-Z]\.?$', part):
                    author = self._parse_single_author(part)
                    if author and author.last_name:
                        authors.append(author)
            return authors

        pattern = r'([A-Z]\.\s+[A-Z][a-z]+)'
        matches = re.findall(pattern, author_text)
        if matches:
            for match in matches:
                author = self._parse_single_author(match.strip())
                if author and author.last_name:
                    authors.append(author)
            return authors

        author = self._parse_single_author(author_text)
        if author and author.last_name:
            authors.append(author)

        return authors

    def _parse_single_author(self, author_str):
        if not author_str:
            return None

        author_str = author_str.strip().rstrip(',')

        hyphen_match = re.search(r'([A-Z])\.-([A-Z])\.', author_str)
        if hyphen_match:
            if ',' in author_str:
                parts = [p.strip() for p in author_str.split(',', 1)]
                if len(parts) == 2:
                    last_name = parts[0]
                    hyphen_initials = re.search(r'([A-Z])\.-([A-Z])\.', parts[1])
                    if hyphen_initials:
                        return Author(last_name=last_name,
                                     first_name=f"{hyphen_initials.group(1)}-{hyphen_initials.group(2)}")

            match = re.match(r'^([A-Z])\.-([A-Z])\.\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)$', author_str)
            if match:
                return Author(last_name=match.group(3),
                             first_name=f"{match.group(1)}-{match.group(2)}")

        if ',' in author_str:
            parts = [p.strip() for p in author_str.split(',', 1)]
            if len(parts) == 2:
                last_name = parts[0]
                initials = parts[1]
                initials_clean = re.sub(r'[.\s]', '', initials)

                if '-' in initials_clean:
                    return Author(last_name=last_name, first_name=initials_clean)

                if '.' in initials and len(initials_clean) >= 2:
                    if len(initials_clean) == 2:
                        return Author(last_name=last_name,
                                     first_name=initials_clean[0],
                                     middle_name=initials_clean[1])
                    elif len(initials_clean) > 2:
                        if len(initials_clean) <= 4 and all(c.isupper() for c in initials_clean):
                            return Author(last_name=last_name,
                                         first_name=initials_clean[0],
                                         middle_name=initials_clean[-1])
                        else:
                            return Author(last_name=last_name, first_name=initials_clean)

                if len(initials_clean) >= 2:
                    return Author(last_name=last_name,
                                 first_name=initials_clean[0],
                                 middle_name=initials_clean[1])
                elif len(initials_clean) == 1:
                    return Author(last_name=last_name,
                                 first_name=initials_clean[0])

                return Author(last_name=last_name)

        match = re.match(r'^([A-Z])\.?\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)$', author_str)
        if match:
            return Author(last_name=match.group(2),
                         first_name=match.group(1))

        match = re.match(r'^([A-Z])\.-([A-Z])\.\s+([A-Z][a-z]+)$', author_str)
        if match:
            return Author(last_name=match.group(3),
                         first_name=f"{match.group(1)}-{match.group(2)}")

        match = re.match(r'^([A-Z])\.\s*([A-Z])\.?\s+([A-Z][a-z]+)$', author_str)
        if match:
            return Author(last_name=match.group(3),
                         first_name=match.group(1),
                         middle_name=match.group(2))

        match = re.match(r'^([A-Z][a-z]+)\s+([A-Z])\.?\s*([A-Z])?\.?$', author_str)
        if match:
            return Author(last_name=match.group(1),
                         first_name=match.group(2),
                         middle_name=match.group(3) if match.group(3) else "")

        match = re.match(r'^([A-Z][a-z]+)\s+([A-Z])\.-([A-Z])\.$', author_str)
        if match:
            return Author(last_name=match.group(1),
                         first_name=f"{match.group(2)}-{match.group(3)}")

        if re.match(r'^[A-Z][a-z]+$', author_str):
            return Author(last_name=author_str)

        return None

    def _extract_url(self, text):
        if not text:
            return ""

        candidates = re.findall(
            r'(https?://[^\s<>"\']+|www\.[^\s<>"\']+)',
            text,
            flags=re.IGNORECASE
        )
        if not candidates:
            return ""

        url = candidates[-1]
        url = url.rstrip(".,;:!?)[]}>\"'")

        if url.lower().startswith('www.'):
            url = 'https://' + url

        return url

    def _strip_urls(self, text):
        if not text:
            return ""
        text = re.sub(r'(https?://[^\s<>"\']+|www\.[^\s<>"\']+)', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def _strip_doi_from_text(self, text, doi=""):
        if not text:
            return ""
        text = re.sub(r'\bdoi\b\s*[:]?\s*10\.[^\s,;]+', ' ', text, flags=re.IGNORECASE)
        if doi:
            d = re.escape(doi)
            text = re.sub(rf'(?:\bdoi\b\s*[:]?\s*)?{d}', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def _extract_doi(self, text):
        if not text:
            return ""

        m = re.search(r'\bdoi\b\s*[:]?\s*([^\s,;]+)', text, re.IGNORECASE)
        if m:
            return self._clean_extracted_value(m.group(1), field_name="doi")

        m = re.search(r'https?://doi\.org/([^\s,;]+)', text, re.IGNORECASE)
        if m:
            return self._clean_extracted_value(m.group(1), field_name="doi")

        try:
            from urllib.parse import urlparse, parse_qs, unquote
            for u in re.findall(r'(https?://[^\s<>"\']+|www\.[^\s<>"\']+)', text, flags=re.IGNORECASE):
                if u.lower().startswith('www.'):
                    u = 'https://' + u
                qs = parse_qs(urlparse(u).query)
                if 'doi' in qs and qs['doi']:
                    return self._clean_extracted_value(unquote(qs['doi'][0]), field_name="doi")
        except Exception:
            pass

        return ""

    def _extract_accessed_date(self, text):
        if not text:
            return ""
        patterns = [
            r'\b(?:accessed|retrieved)\b\s*(?:on\s*)?([A-Za-z]+\s+\d{1,2},\s*\d{4})',
            r'\b(?:accessed|retrieved)\b\s*(?:on\s*)?(\d{4}-\d{2}-\d{2})',
            r'\b(?:accessed|retrieved)\b\s*(?:on\s*)?(\d{1,2}\.\d{1,2}\.\d{4})',
            r'\[\s*(?:accessed|retrieved)\s*[:]?\s*([^\]]+)\]',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                return self._clean_extracted_value(m.group(1), field_name="accessed_date")
        return ""

    def _extract_publisher(self, text):
        if not text:
            return ""
        m = re.search(r'\bpublisher\s*[:]?\s*([^.,;]+)', text, re.IGNORECASE)
        if m:
            return self._clean_extracted_value(m.group(1), field_name="publisher")
        m = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*:\s*([^.,;]+)', text)
        if m:
            return self._clean_extracted_value(m.group(2), field_name="publisher")
        return ""

    def _extract_city_and_publisher(self, text):
        if not text:
            return "", ""
        m = re.search(r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*:\s*([^.,;]+)', text)
        if m:
            return m.group(1).strip(), m.group(2).strip()
        return "", self._extract_publisher(text)

    def _extract_isbn(self, text):
        if not text:
            return ""
        m = re.search(r'\bISBN\b\s*[:]?\s*([0-9Xx\-]+)', text)
        return self._clean_extracted_value(m.group(1), field_name="isbn") if m else ""

    def _extract_edition(self, text):
        if not text:
            return ""
        patterns = [
            r'\b(\d+(?:st|nd|rd|th)\s+ed\.?)(?!\w)',
            r'\b(\d+\s*ed(?:ition)?\.?)(?!\w)',
            r'\b(ed\.?\s*\d+)\b',
            r'\b(rev\.?\s*ed\.?)\b',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                return self._clean_extracted_value(m.group(1), field_name="edition")
        return ""

    def _extract_website(self, text, url=""):
        if not text:
            return ""
        m = re.search(r"\b(?:on|at)\s+([^.,;]+)\s*(?:website|site)\b", text, re.IGNORECASE)
        if m:
            return self._clean_extracted_value(m.group(1), field_name="website")
        return ""

    def _parse_article(self, text):
        item = Article()

        authors, title, rest = self._extract_authors_and_title(text)
        item.authors = authors
        item.title = title

        item.url = self._clean_extracted_value(self._extract_url(text) or self._extract_url(rest), field_name="url")
        item.doi = self._clean_extracted_value(self._extract_doi(text) or self._extract_doi(rest), field_name="doi")
        item.accessed_date = self._clean_extracted_value(self._extract_accessed_date(text) or self._extract_accessed_date(rest), field_name="accessed_date")

        rest = self._strip_urls(rest)
        rest = self._strip_doi_from_text(rest, item.doi)

        year_match = re.search(r'\b(19|20)\d{2}\b', rest)
        if year_match:
            item.year = int(year_match.group())

        journal_patterns = [
            r'(IEEE\s+Trans\.\s+on\s+[^.,]+)',
            r'(IEEE\s+Trans\.\s+[^.,]+)',
            r'\.\s*([A-Z][^.,]*(?:\s+[A-Z][^.,]*)*)(?=\s*[,\\.])',
        ]

        for pattern in journal_patterns:
            match = re.search(pattern, rest)
            if match:
                journal = match.group(1).strip()
                if len(journal) > 3 and not journal.isdigit():
                    item.journal = self._clean_extracted_value(journal, field_name="journal")
                    break

        if not item.journal:
            parts = [p.strip() for p in rest.split('.')]
            for part in parts:
                if (len(part) > 3 and re.match(r'^[A-Z]', part) and
                    not re.search(r'\b(vol|no|pp|p|doi|http)\b', part, re.I) and
                    not re.search(r'\d{4}', part)):
                    item.journal = self._clean_extracted_value(part, field_name="journal")
                    break

        vol_patterns = [
            r'\b[Vv]olume\b\s*[:]?\s*(\d+|[IVXLCDM]+)',
            r'[Vv]ol\.?\s*[:]?\s*(\d+|[IVXLCDM]+)',
            r',\s*(\d+)\s*,',
            r'\.\s*(\d+)\s*,',
        ]
        for pattern in vol_patterns:
            vol_match = re.search(pattern, rest)
            if vol_match:
                item.volume = self._clean_extracted_value(vol_match.group(1), field_name="volume")
                break

        issue_patterns = [
            r'\b(?:no\.|no|issue)\b\s*[:]?\s*(\d+)\b',
            r'\b\d+\s*\(\s*(\d+)\s*\)\b',
        ]
        for pattern in issue_patterns:
            issue_match = re.search(pattern, rest)
            if issue_match:
                item.issue = self._clean_extracted_value(issue_match.group(1), field_name="issue")
                break

        pages_patterns = [
            r'[Pp]p?\.?\s*[:]?\s*(\d+[-–]\d+)',
            r'\b(\d+)\s*[-–]\s*(\d+)\b',
            r',\s*(\d+[-–]\d+)\s*\.',
            r',\s*(\d+[-–]\d+)\s*,',
        ]
        for pattern in pages_patterns:
            pages_match = re.search(pattern, rest)
            if pages_match:
                if len(pages_match.groups()) == 2:
                    item.pages = self._clean_extracted_value(f"{pages_match.group(1)}-{pages_match.group(2)}", field_name="pages")
                else:
                    item.pages = self._clean_extracted_value(pages_match.group(1), field_name="pages")
                break

        pub = self._extract_publisher(rest) or self._extract_publisher(text)
        if pub:
            item.publisher = self._clean_extracted_value(pub, field_name="publisher")

        return item

    def _parse_book(self, text):
        item = Book()

        authors, title, rest = self._extract_authors_and_title(text)
        item.authors = authors
        item.title = title

        item.url = self._clean_extracted_value(self._extract_url(text) or self._extract_url(rest), field_name="url")
        item.doi = self._clean_extracted_value(self._extract_doi(text) or self._extract_doi(rest), field_name="doi")
        item.accessed_date = self._clean_extracted_value(self._extract_accessed_date(text) or self._extract_accessed_date(rest), field_name="accessed_date")

        year_match = re.search(r'\b(19|20)\d{2}\b', rest)
        if year_match:
            item.year = int(year_match.group())

        city, publisher = self._extract_city_and_publisher(rest)
        if city and not item.city:
            item.city = self._clean_extracted_value(city, field_name="city")
        if publisher and not item.publisher:
            item.publisher = self._clean_extracted_value(publisher, field_name="publisher")

        if not item.publisher:
            item.publisher = self._clean_extracted_value(self._extract_publisher(rest), field_name="publisher")

        item.edition = self._clean_extracted_value(self._extract_edition(rest), field_name="edition")

        item.isbn = self._clean_extracted_value(self._extract_isbn(rest), field_name="isbn")

        return item

    def _parse_conference(self, text):
        item = ConferencePaper()

        authors, title, rest = self._extract_authors_and_title(text)
        item.authors = authors
        item.title = title

        item.url = self._clean_extracted_value(self._extract_url(text) or self._extract_url(rest), field_name="url")
        item.doi = self._clean_extracted_value(self._extract_doi(text) or self._extract_doi(rest), field_name="doi")
        item.accessed_date = self._clean_extracted_value(self._extract_accessed_date(text) or self._extract_accessed_date(rest), field_name="accessed_date")

        year_match = re.search(r'\b(19|20)\d{2}\b', rest)
        if year_match:
            item.year = int(year_match.group())

        conf_match = re.search(r'\b[Ii]n\s*:?\s*([^.,]+)', rest)
        if conf_match:
            item.conference_name = self._clean_extracted_value(conf_match.group(1).strip(), field_name="conference_name")

        pages_patterns = [
            r'[Pp]p?\.?\s*[:]?\s*(\d+[-–]\d+)',
            r'\b(\d+)\s*[-–]\s*(\d+)\b',
            r',\s*(\d+[-–]\d+)\s*\.',
            r',\s*(\d+[-–]\d+)\s*,',
        ]
        for pattern in pages_patterns:
            m = re.search(pattern, rest)
            if m:
                if len(m.groups()) == 2:
                    item.pages = self._clean_extracted_value(f"{m.group(1)}-{m.group(2)}", field_name="pages")
                else:
                    item.pages = self._clean_extracted_value(m.group(1), field_name="pages")
                break

        loc_match = re.search(r'\(([^\)]+)\)', rest)
        if loc_match and not re.search(r'\d', loc_match.group(1)):
            item.location = self._clean_extracted_value(loc_match.group(1).strip(), field_name="location")
        if not item.location:
            loc_match = re.search(r',\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*?(?:,\s*[A-Z][a-z]+)?)\s*[\.,]', rest)
            if loc_match:
                item.location = self._clean_extracted_value(loc_match.group(1).strip(), field_name="location")

        pub_match = re.search(r'\b(IEEE|ACM|Springer|Elsevier)\b', rest)
        if pub_match:
            item.publisher = self._clean_extracted_value(pub_match.group(1), field_name="publisher")

        return item

    def _parse_electronic_resource(self, text):
        item = ElectronicResource()

        authors, title, rest = self._extract_authors_and_title(text)
        item.authors = authors
        item.title = title

        if (not item.authors or not item.title) and text:
            parts = [p.strip() for p in text.split('.') if p.strip()]
            if len(parts) >= 2:
                potential_authors = parts[0]
                potential_title = parts[1]
                if not item.authors:
                    item.authors = self._parse_authors(potential_authors)
                if not item.title:
                    item.title = potential_title

        item.url = self._clean_extracted_value(self._extract_url(text) or self._extract_url(rest), field_name="url")
        item.doi = self._clean_extracted_value(self._extract_doi(text) or self._extract_doi(rest), field_name="doi")
        item.accessed_date = self._clean_extracted_value(self._extract_accessed_date(text) or self._extract_accessed_date(rest), field_name="accessed_date")

        year_match = re.search(r'\b(19|20)\d{2}\b', text)
        if year_match:
            item.year = int(year_match.group())

        item.website = self._clean_extracted_value(self._extract_website(text, item.url), field_name="website")

        pub = self._extract_publisher(text)
        if pub:
            item.publisher = self._clean_extracted_value(pub, field_name="publisher")

        return item

    def format_all_items(self):
        if not self.current_style:
            raise ValueError("Стиль не установлен")

        return [self.current_style.format_item(item) for item in self.items]

    def validate_all_items(self):
        if not self.current_style:
            raise ValueError("Стиль не установлен")

        results = []
        for item in self.items:
            is_valid, missing = self.current_style.validate_item(item)
            results.append((item, is_valid, missing))

        return results

    def save_to_docx(self, filepath, highlight_missing=True):
        doc = Document()
        doc.add_heading('Список литературы', 0)

        if not self.current_style:
            doc.add_paragraph("Стиль не установлен")
            doc.save(filepath)
            return

        validation_results = []
        if highlight_missing:
            validation_results = self.validate_all_items()

        for i, item in enumerate(self.items, 1):
            formatted = self.current_style.format_item(item)
            para = doc.add_paragraph()

            run = para.add_run(f"{i}. ")
            run.bold = True

            para.add_run(formatted)

            if highlight_missing and validation_results:
                _, is_valid, missing = validation_results[i-1]
                if not is_valid and missing:
                    warning = f" [Отсутствуют: {', '.join(missing)}]"
                    run = para.add_run(warning)
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.italic = True

        doc.save(filepath)

    def create_custom_style(self, field_order, required_fields,
                          author_config=None):
        style = CitationStyle("Пользовательский")
        style.set_field_order(field_order)
        style.set_required_fields(required_fields)

        if author_config:
            style.set_author_format(author_config)

        return style